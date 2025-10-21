import csv
import os
from docx import Document

# === Configuration ===
csv_file = r"C:\Users\ahmad.satria.psi\Documents\Telkom\Daftar_Hadir_UT\source.csv"
template_path = r"C:\Users\ahmad.satria.psi\Documents\Telkom\Daftar_Hadir_UT\Master_Daftar Hadir UT.docx"
output_folder = r"C:\Users\ahmad.satria.psi\Documents\Telkom\Daftar_Hadir_UT\output"

# Create folder if it doesn’t exist
os.makedirs(output_folder, exist_ok=True)

def replace_text_in_paragraphs(doc, replacements):
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(key, value)

def replace_text_in_tables(doc, replacements):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.text = run.text.replace(key, value)

# === Read CSV and process ===
with open(csv_file, encoding="utf-8-sig") as f:
    reader = csv.DictReader(f, delimiter=';')
    for row in reader:
        site_id = row['site_id'].strip()
        site_name = row['site_name'].strip()
        
        # Load the Word template
        doc = Document(template_path)
        
        # Define what to replace
        replacements = {
            '%site_id%': site_id,
            '%site_name%': site_name
        }
        
        # Replace inside paragraphs and tables
        #replace_text_in_paragraphs(doc, replacements)
        replace_text_in_tables(doc, replacements)
        
        # Make safe filename
        safe_name = "".join(c if c.isalnum() or c in " -_." else "_" for c in site_name)
        output_path = os.path.join(output_folder, f"{safe_name}.docx")
        
        # Save file
        doc.save(output_path)
        print(f"Saved: {output_path}")

print("✅ All files generated successfully!")
