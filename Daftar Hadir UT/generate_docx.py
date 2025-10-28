import csv
import os
from docx import Document
from docx2pdf import convert

# === Configuration ===
csv_file = r"source.csv"
template_path = r"Master_Daftar Hadir UT.docx"
output_folder = r"output"

os.makedirs(output_folder, exist_ok=True)


def replace_text_in_paragraph(paragraph, replacements):
    """Replace placeholders inside a single paragraph, preserving runs."""
    full_text = ""
    for run in paragraph.runs:
        full_text += run.text

    for key, value in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, value)

    # Rewrite runs
    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(full_text)


def replace_text_in_table(table, replacements):
    """Replace text inside all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)


def process_document(template_path, output_folder, site_id, site_name):
    """Open template, replace placeholders, save new file."""
    doc = Document(template_path)

    replacements = {
        "%site_id%": site_id,
        "%site_name%": site_name
    }

    # Replace in paragraphs
    # for paragraph in doc.paragraphs:
    #     replace_text_in_paragraph(paragraph, replacements)

    # Replace inside first table found
    replace_text_in_table(doc.tables[0], replacements)

    safe_name = "".join(c if c.isalnum() or c in " -_." else "_" for c in site_name)
    output_path_docx = os.path.join(output_folder, f"{safe_name}.docx")
    doc.save(output_path_docx)
    print(f"✅ Saved docx: {output_path_docx}")


# === MAIN LOOP ===
with open(csv_file, encoding="utf-8-sig") as f:
    reader = csv.DictReader(f, delimiter=';')
    for row in reader:
        site_id = row['site_id'].strip()
        site_name = row['site_name'].strip()
        process_document(template_path, output_folder, site_id, site_name)

print("✅ Convert all docx file into pdf")
convert(r"output")
print("🎉 All files generated successfully!")
